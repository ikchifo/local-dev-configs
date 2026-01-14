#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.9"
# dependencies = [
#     "python-docx",
#     "matplotlib",
# ]
# ///
"""
RFC to DOCX Converter

Converts markdown-like RFC text (including ASCII diagrams) to a properly formatted DOCX document.

Usage:
    python rfc_to_docx.py input.md output.docx

Or as a module:
    from rfc_to_docx import convert_rfc_to_docx
    convert_rfc_to_docx(markdown_text, "output.docx")
"""

import argparse
import io
import re
import sys
from pathlib import Path
from typing import List, Tuple, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch


# Font configuration
DEFAULT_FONT = "Proxima Nova"
CODE_FONT = "Courier New"
TEXT_COLOR = RGBColor(0, 0, 0)


def set_run_font(run, font_name=DEFAULT_FONT, color=TEXT_COLOR, size=None):
    """Set font name, color, and optionally size for a run."""
    run.font.color.rgb = color
    run.font.name = font_name
    if size:
        run.font.size = Pt(size)

    # Set the font for complex scripts as well
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def add_heading(doc, text: str, level: int):
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_run_font(run, DEFAULT_FONT)
    return heading


def add_paragraph(doc, text: str):
    """Add a paragraph with proper formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, DEFAULT_FONT)
    return para


def add_bullet_point(doc, text: str, level: int = 0):
    """Add a bullet point with proper formatting."""
    para = doc.add_paragraph(style="List Bullet")
    para.clear()
    run = para.add_run(text)
    set_run_font(run, DEFAULT_FONT)
    if level > 0:
        para.paragraph_format.left_indent = Inches(0.25 * level)
    return para


def add_numbered_item(doc, text: str, number: int):
    """Add a numbered item with proper formatting."""
    para = doc.add_paragraph()
    run = para.add_run(f"{number}. {text}")
    set_run_font(run, DEFAULT_FONT)
    return para


def add_code_block(doc, code: str):
    """Add a code block with monospace formatting."""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = CODE_FONT
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_COLOR

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), CODE_FONT)
    rFonts.set(qn('w:hAnsi'), CODE_FONT)
    rFonts.set(qn('w:cs'), CODE_FONT)

    para.paragraph_format.left_indent = Inches(0.25)

    # Add light gray background shading
    from docx.oxml import OxmlElement
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    para._element.get_or_add_pPr().append(shd)

    return para


def add_table(doc, headers: List[str], rows: List[List[str]]):
    """Add a table with proper formatting."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for run in header_cells[i].paragraphs[0].runs:
            run.bold = True
            set_run_font(run, DEFAULT_FONT)

    # Add rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)
            for run in row_cells[col_idx].paragraphs[0].runs:
                set_run_font(run, DEFAULT_FONT)

    return table


def is_ascii_diagram(text: str) -> bool:
    """Detect if a code block contains an ASCII diagram."""
    diagram_chars = ['─', '│', '┌', '┐', '└', '┘', '├', '┤', '┬', '┴', '┼',
                     '═', '║', '╔', '╗', '╚', '╝', '╠', '╣', '╦', '╩', '╬',
                     '+', '-', '|', '*', '/', '\\', '<', '>', '^', 'v',
                     '┌─', '──', '│ ', '└─', '├─', '─┤', '─┐', '─┘']

    # Check for box-drawing patterns
    box_patterns = [
        r'[\+\-\|]{3,}',  # +---+ patterns
        r'[┌┐└┘├┤┬┴┼─│]+',  # Unicode box drawing
        r'\[.*\].*\[.*\]',  # [Box] -> [Box] patterns
        r'─{3,}',  # Long horizontal lines
        r'│\s*│',  # Vertical bars
    ]

    for pattern in box_patterns:
        if re.search(pattern, text):
            return True

    # Check character density of diagram characters
    diagram_char_count = sum(1 for c in text if c in '─│┌┐└┘├┤┬┴┼═║╔╗╚╝╠╣╦╩╬+-|*/\\<>^v[](){}')
    total_chars = len(text.replace(' ', '').replace('\n', ''))

    if total_chars > 0 and diagram_char_count / total_chars > 0.15:
        return True

    return False


def ascii_diagram_to_image(diagram_text: str) -> io.BytesIO:
    """Convert ASCII diagram to an image using matplotlib."""
    lines = diagram_text.strip().split('\n')

    # Calculate figure size based on content
    max_width = max(len(line) for line in lines) if lines else 40
    height = len(lines)

    fig_width = min(12, max(6, max_width * 0.12))
    fig_height = min(10, max(2, height * 0.25))

    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    ax.set_xlim(0, max_width)
    ax.set_ylim(0, height)
    ax.axis('off')

    # Set font
    plt.rcParams['font.family'] = 'monospace'

    # Render each line
    for i, line in enumerate(lines):
        y = height - i - 0.5
        ax.text(0, y, line, fontsize=9, fontfamily='monospace',
                verticalalignment='center', color='black')

    # Add a subtle border
    border = FancyBboxPatch((0, 0), max_width, height,
                             boxstyle="round,pad=0.02",
                             facecolor='#FAFAFA',
                             edgecolor='#DDDDDD',
                             linewidth=1,
                             zorder=-1)
    ax.add_patch(border)

    plt.tight_layout(pad=0.5)

    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close()

    return buf


def parse_markdown_table(lines: List[str]) -> Tuple[List[str], List[List[str]]]:
    """Parse a markdown table into headers and rows."""
    headers = []
    rows = []

    for i, line in enumerate(lines):
        line = line.strip()
        if not line.startswith('|'):
            continue

        cells = [cell.strip() for cell in line.split('|')[1:-1]]

        # Skip separator line (contains only dashes)
        if all(re.match(r'^[-:]+$', cell) for cell in cells):
            continue

        if not headers:
            headers = cells
        else:
            rows.append(cells)

    return headers, rows


def parse_rfc_content(text: str) -> List[Tuple[str, any]]:
    """
    Parse RFC markdown-like content into structured elements.

    Returns list of tuples: (element_type, content)
    Element types: 'title', 'h1', 'h2', 'h3', 'paragraph', 'bullet',
                   'numbered', 'code', 'diagram', 'table', 'metadata'
    """
    elements = []
    lines = text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Title (# at start of document or ## RFC:)
        if stripped.startswith('# ') and not elements:
            elements.append(('title', stripped[2:].strip()))
            i += 1
            continue

        # Headings
        if stripped.startswith('### '):
            elements.append(('h3', stripped[4:].strip()))
            i += 1
            continue
        elif stripped.startswith('## '):
            elements.append(('h2', stripped[3:].strip()))
            i += 1
            continue
        elif stripped.startswith('# '):
            elements.append(('h1', stripped[2:].strip()))
            i += 1
            continue

        # Code blocks
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```

            code_content = '\n'.join(code_lines)

            # Check if it's an ASCII diagram
            if is_ascii_diagram(code_content):
                elements.append(('diagram', code_content))
            else:
                elements.append(('code', code_content))
            continue

        # Tables (markdown style)
        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            headers, rows = parse_markdown_table(table_lines)
            if headers:
                elements.append(('table', (headers, rows)))
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            bullet_text = stripped[2:].strip()
            indent_level = (len(line) - len(line.lstrip())) // 2
            elements.append(('bullet', (bullet_text, indent_level)))
            i += 1
            continue

        # Numbered items
        match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if match:
            num = int(match.group(1))
            text = match.group(2)
            elements.append(('numbered', (text, num)))
            i += 1
            continue

        # Metadata lines (Key: Value at start of document)
        if ':' in stripped and len(elements) <= 2:
            key, _, value = stripped.partition(':')
            if key.strip() in ['Author', 'Status', 'Created', 'Date', 'Version']:
                elements.append(('metadata', (key.strip(), value.strip())))
                i += 1
                continue

        # Regular paragraph - collect consecutive non-empty lines
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line or next_line.startswith('#') or next_line.startswith('```') or \
               next_line.startswith('|') or next_line.startswith('- ') or \
               next_line.startswith('* ') or re.match(r'^\d+\.\s', next_line):
                break
            para_lines.append(next_line)
            i += 1

        elements.append(('paragraph', ' '.join(para_lines)))

    return elements


def convert_rfc_to_docx(markdown_text: str, output_path: str, title: str = None):
    """
    Convert RFC markdown text to a DOCX document.

    Args:
        markdown_text: The RFC content in markdown-like format
        output_path: Path to save the DOCX file
        title: Optional override for document title
    """
    doc = Document()
    elements = parse_rfc_content(markdown_text)

    metadata_items = []

    for elem_type, content in elements:
        if elem_type == 'title':
            heading = doc.add_heading(title or content, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                set_run_font(run, DEFAULT_FONT)

        elif elem_type == 'metadata':
            metadata_items.append(content)

        elif elem_type == 'h1':
            # Flush metadata before first h1
            if metadata_items:
                meta = doc.add_paragraph()
                for key, value in metadata_items:
                    run = meta.add_run(f"{key}: ")
                    run.bold = True
                    set_run_font(run, DEFAULT_FONT)
                    run = meta.add_run(f"{value}\n")
                    set_run_font(run, DEFAULT_FONT)
                metadata_items = []
                doc.add_paragraph()

            add_heading(doc, content, 1)

        elif elem_type == 'h2':
            add_heading(doc, content, 2)

        elif elem_type == 'h3':
            add_heading(doc, content, 3)

        elif elem_type == 'paragraph':
            add_paragraph(doc, content)

        elif elem_type == 'bullet':
            text, level = content
            add_bullet_point(doc, text, level)

        elif elem_type == 'numbered':
            text, num = content
            add_numbered_item(doc, text, num)

        elif elem_type == 'code':
            add_code_block(doc, content)

        elif elem_type == 'diagram':
            # Convert ASCII diagram to image
            try:
                img_buf = ascii_diagram_to_image(content)
                doc.add_picture(img_buf, width=Inches(6.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                # Fall back to code block if image conversion fails
                add_code_block(doc, content)

        elif elem_type == 'table':
            headers, rows = content
            add_table(doc, headers, rows)
            doc.add_paragraph()  # Space after table

    # Flush any remaining metadata
    if metadata_items:
        meta = doc.add_paragraph()
        for key, value in metadata_items:
            run = meta.add_run(f"{key}: ")
            run.bold = True
            set_run_font(run, DEFAULT_FONT)
            run = meta.add_run(f"{value}\n")
            set_run_font(run, DEFAULT_FONT)

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description='Convert RFC markdown to DOCX with proper formatting'
    )
    parser.add_argument('input', help='Input markdown file path')
    parser.add_argument('output', help='Output DOCX file path')
    parser.add_argument('--title', help='Override document title')

    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    output_path = convert_rfc_to_docx(content, args.output, args.title)
    print(f"RFC document saved to: {output_path}")


if __name__ == '__main__':
    main()
